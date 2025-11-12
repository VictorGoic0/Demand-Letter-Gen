"""
DOCX generator for converting HTML letter content to .docx format.
Handles HTML parsing and conversion to python-docx Document objects.
"""
import logging
import re
import io
from datetime import datetime
from typing import Optional
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)


class HTMLToDocxParser(HTMLParser):
    """
    HTML parser that converts HTML content to python-docx Document.
    Handles common HTML tags: p, h1, h2, h3, strong, b, em, i, ul, ol, li
    """
    
    def __init__(self, doc: Document):
        """
        Initialize parser with a Document object.
        
        Args:
            doc: python-docx Document object to populate
        """
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.formatting_stack = []  # Stack to track nested formatting (bold, italic)
        self.list_stack = []  # Track nested lists
        
    def _ensure_paragraph(self):
        """Ensure we have a current paragraph, creating one if needed."""
        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph()
    
    def _create_run(self):
        """Create a new run with current formatting from stack."""
        self._ensure_paragraph()
        run = self.current_paragraph.add_run()
        
        # Apply formatting from stack
        for fmt in self.formatting_stack:
            if fmt == 'bold':
                run.bold = True
            elif fmt == 'italic':
                run.italic = True
        
        return run
        
    def handle_starttag(self, tag: str, attrs: list):
        """Handle opening HTML tags."""
        if tag in ['p', 'h1', 'h2', 'h3']:
            # Create new paragraph
            self.current_paragraph = self.doc.add_paragraph()
            
            # Set heading style
            if tag == 'h1':
                self.current_paragraph.style = 'Heading 1'
            elif tag == 'h2':
                self.current_paragraph.style = 'Heading 2'
            elif tag == 'h3':
                self.current_paragraph.style = 'Heading 3'
            # p tags use default paragraph style
            
        elif tag in ['strong', 'b']:
            # Bold formatting - add to stack
            self.formatting_stack.append('bold')
            
        elif tag in ['em', 'i']:
            # Italic formatting - add to stack
            self.formatting_stack.append('italic')
            
        elif tag == 'ul':
            # Start unordered list
            self.list_stack.append('ul')
            
        elif tag == 'ol':
            # Start ordered list
            self.list_stack.append('ol')
            
        elif tag == 'li':
            # List item
            if not self.list_stack:
                # If no list context, treat as regular paragraph
                self.current_paragraph = self.doc.add_paragraph()
            else:
                # Create paragraph with list formatting
                list_style = 'List Bullet' if self.list_stack[-1] == 'ul' else 'List Number'
                self.current_paragraph = self.doc.add_paragraph(style=list_style)
            
    def handle_endtag(self, tag: str):
        """Handle closing HTML tags."""
        if tag in ['p', 'h1', 'h2', 'h3']:
            # End paragraph
            self.current_paragraph = None
            
        elif tag in ['strong', 'b']:
            # End bold formatting - pop from stack (LIFO)
            # Remove the last occurrence (most recent)
            if 'bold' in self.formatting_stack:
                # Reverse to find last occurrence, then remove it
                idx = len(self.formatting_stack) - 1 - self.formatting_stack[::-1].index('bold')
                self.formatting_stack.pop(idx)
                    
        elif tag in ['em', 'i']:
            # End italic formatting - pop from stack (LIFO)
            # Remove the last occurrence (most recent)
            if 'italic' in self.formatting_stack:
                # Reverse to find last occurrence, then remove it
                idx = len(self.formatting_stack) - 1 - self.formatting_stack[::-1].index('italic')
                self.formatting_stack.pop(idx)
                    
        elif tag == 'ul':
            # End unordered list
            if self.list_stack and self.list_stack[-1] == 'ul':
                self.list_stack.pop()
                    
        elif tag == 'ol':
            # End ordered list
            if self.list_stack and self.list_stack[-1] == 'ol':
                self.list_stack.pop()
                    
        elif tag == 'li':
            # End list item
            self.current_paragraph = None
            
    def handle_data(self, data: str):
        """Handle text content."""
        # Don't strip all whitespace - preserve some structure
        # But remove leading/trailing whitespace on lines
        lines = data.split('\n')
        text_parts = []
        for line in lines:
            stripped = line.strip()
            if stripped:
                text_parts.append(stripped)
            elif text_parts:  # Preserve single newline between content
                text_parts.append(' ')
        
        text = ''.join(text_parts).strip()
        if not text:
            return
            
        # Create run with current formatting
        run = self._create_run()
        run.add_text(text)


def html_to_docx(html_content: str) -> Document:
    """
    Convert HTML content to a python-docx Document object.
    
    Args:
        html_content: HTML string to convert
        
    Returns:
        Document object with converted content
        
    Raises:
        ValueError: If HTML parsing fails
    """
    try:
        # Create new document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Parse HTML
        parser = HTMLToDocxParser(doc)
        
        # Clean HTML - remove script and style tags
        cleaned_html = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        cleaned_html = re.sub(r'<style[^>]*>.*?</style>', '', cleaned_html, flags=re.DOTALL | re.IGNORECASE)
        
        # Parse HTML
        parser.feed(cleaned_html)
        
        logger.info("Successfully converted HTML to DOCX")
        return doc
        
    except Exception as e:
        logger.error(f"Failed to convert HTML to DOCX: {str(e)}")
        raise ValueError(f"Failed to convert HTML to DOCX: {str(e)}") from e


def generate_filename(letter_title: str, date: Optional[datetime] = None) -> str:
    """
    Generate a sanitized filename for the DOCX file.
    
    Format: Demand_Letter_[Title]_[Date].docx
    
    Args:
        letter_title: Letter title to use in filename
        date: Optional date to include (defaults to current date)
        
    Returns:
        Sanitized filename (max 50 chars)
    """
    try:
        # Use provided date or current date
        if date is None:
            date = datetime.utcnow()
        
        # Format date as YYYY-MM-DD
        date_str = date.strftime("%Y-%m-%d")
        
        # Sanitize title: keep alphanumeric, spaces, underscores, hyphens
        # Remove all other special characters
        sanitized_title = re.sub(r'[^a-zA-Z0-9\s_-]', '', letter_title)
        
        # Replace spaces with underscores
        sanitized_title = re.sub(r'\s+', '_', sanitized_title)
        
        # Build filename
        filename = f"Demand_Letter_{sanitized_title}_{date_str}.docx"
        
        # Truncate to 50 characters if needed (keep .docx extension)
        if len(filename) > 50:
            # Leave room for .docx (5 chars) and date (11 chars) and "Demand_Letter_" (14 chars)
            max_title_length = 50 - 14 - 11 - 5  # = 20
            if max_title_length > 0:
                truncated_title = sanitized_title[:max_title_length]
                filename = f"Demand_Letter_{truncated_title}_{date_str}.docx"
            else:
                # Fallback: just use date
                filename = f"Demand_Letter_{date_str}.docx"
        
        logger.info(f"Generated filename: {filename}")
        return filename
        
    except Exception as e:
        logger.error(f"Failed to generate filename: {str(e)}")
        # Fallback filename
        date_str = (date or datetime.utcnow()).strftime("%Y-%m-%d")
        return f"Demand_Letter_{date_str}.docx"


def save_docx_to_s3(
    doc: Document,
    bucket_name: str,
    s3_key: str,
    s3_client,
) -> str:
    """
    Save a Document object to S3 as a .docx file.
    
    Args:
        doc: python-docx Document object
        bucket_name: S3 bucket name
        s3_key: S3 key (path) for the file
        s3_client: S3Client instance
        
    Returns:
        S3 key where file was saved
        
    Raises:
        S3UploadException: If upload fails
    """
    try:
        # Save document to BytesIO buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Upload to S3
        s3_client.upload_fileobj(
            file_obj=buffer,
            bucket_name=bucket_name,
            s3_key=s3_key,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        
        logger.info(f"DOCX file saved to S3: s3://{bucket_name}/{s3_key}")
        return s3_key
        
    except Exception as e:
        logger.error(f"Failed to save DOCX to S3: {str(e)}")
        from shared.exceptions import S3UploadException
        raise S3UploadException(
            message="Failed to upload DOCX file to S3",
            detail=str(e),
        ) from e

